from flask import Flask, render_template, request, redirect, send_file, send_from_directory
import os
import pdf2docx
from werkzeug.utils import secure_filename
from docx import Document
from docx2pdf import convert
from PIL import Image
from fpdf import FPDF
from pdf2image import convert_from_path
import PyPDF2
import pandas as pd
import fpdf
from docx import Document



app = Flask(__name__)

# Set the upload folder path
UPLOAD_FOLDER = 'uploads'

app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER
ALLOWED_EXTENSIONS = {'xlsx', 'xls','docx'}

# Route to handle the upload page


@app.route('/')
def home():
    return render_template('home.html')

@app.route('/upload')
def index():
    return render_template('index.html')


@app.route('/word')
def word():
    return render_template('word.html')


@app.route('/jpg')
def jpg():
    return render_template('jpg.html')


@app.route('/pdf')
def pdf():
    return render_template('pdf.html')


def convert_pdf_to_excel(file_path, output_path):
    with open(file_path, 'rb') as pdf_file:
        pdf_reader = PyPDF2.PdfReader(pdf_file)
        num_pages = len(pdf_reader.pages)
        data = []
        for page_num in range(num_pages):
            page = pdf_reader.pages[page_num]
            page_data = page.extract_text()
            data.append(page_data)

        df = pd.DataFrame(data)
        df.to_excel(output_path, index=False)

@app.route('/pdfe')
def pdfe():
    return render_template('pdfe.html')

@app.route('/excel')
def excel():
    return render_template('excel.html')


@app.route('/jpgw')
def jpgw():
    return render_template('jpgw.html')

# Route to handle file upload and conversion


@app.route('/upload', methods=['POST'])
def upload_file():
    file = request.files['file']

    filename = secure_filename(file.filename)
    file.save(os.path.join(app.config['UPLOAD_FOLDER'], filename))
    pdf_file_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
    docx_file_path = os.path.join(
        app.config['UPLOAD_FOLDER'], f"{os.path.splitext(filename)[0]}.docx")
    pdf2docx.parse(pdf_file_path, docx_file_path)
    return send_file(docx_file_path, as_attachment=True)


@app.route('/', methods=['POST'])
def contact():
    # Get the uploaded file
    file = request.files['file']

    # Save the file to the uploads folder
    file.save(f"{app.config['UPLOAD_FOLDER']}/{file.filename}")

    # Convert the Word document to PDF
    doc = Document(f"{app.config['UPLOAD_FOLDER']}/{file.filename}")
    pdf = FPDF()
    
    for para in doc.paragraphs:
        pdf.add_page()
        pdf.set_font("Arial", size=12)
        pdf.cell(0, 10, txt=para.text, ln=1)

    pdf.output(f"{app.config['UPLOAD_FOLDER']}/{file.filename.split('.')[0]}.pdf")

    return render_template('download_word.html', filename=file.filename)



@app.route('/convert', methods=['POST'])
def convert():
    # Get the uploaded image file
    image_file = request.files['image']

    # Save the image to the upload folder
    image_path = f"{app.config['UPLOAD_FOLDER']}/{image_file.filename}"
    image_file.save(image_path)

    # Convert the image to PDF
    pdf_path = f"{app.config['UPLOAD_FOLDER']}/{image_file.filename.split('.')[0]}.pdf"
    image = Image.open(image_path)
    pdf = FPDF()
    pdf.add_page()
    pdf.image(image_path, 0, 0, pdf.w, pdf.h)
    pdf.output(pdf_path, "F")

    # Provide the download link for the converted PDF file
    return render_template('download_jpg.html', pdf_filename=f"{image_file.filename.split('.')[0]}.pdf")


@app.route('/transfer', methods=['POST'])
def transfer():
    # Get the uploaded PDF file
    file = request.files['pdf']

    # Save the PDF file
    pdf_path = os.path.join(app.config['UPLOAD_FOLDER'], file.filename)
    file.save(pdf_path)

    # Convert PDF to JPG
    images = convert_from_path(pdf_path)

    # Save each image as JPG
    jpg_filenames = []
    for i, image in enumerate(images):
        jpg_filename = f"{i}.jpg"
        jpg_path = os.path.join(app.config['UPLOAD_FOLDER'], jpg_filename)
        image.save(jpg_path)
        jpg_filenames.append(jpg_filename)

    return render_template('download_pdf_jpg.html', filenames=jpg_filenames)


@app.route('/converted', methods=['POST'])
def converted():
    file = request.files['pdf_file']
    filename = file.filename
    file_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
    file.save(file_path)

    output_file = os.path.splitext(filename)[0] + '.xlsx'
    output_path = os.path.join(app.config['UPLOAD_FOLDER'], output_file)

    convert_pdf_to_excel(file_path, output_path)

    return render_template('download_pdf_excel.html', filename=output_file)


@app.route('/update', methods=['POST'])
def update():
    file = request.files['file']
    if file and allowed_file(file.filename):
        filename = file.filename
        file_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
        file.save(file_path)

        output_filename = os.path.splitext(filename)[0] + '.pdf'
        output_path = os.path.join(app.config['UPLOAD_FOLDER'], output_filename)

        # Read Excel file using pandas
        df = pd.read_excel(file_path)

        # Generate PDF from DataFrame
        pdf = fpdf.FPDF()
        pdf.add_page()
        pdf.set_auto_page_break(auto=True, margin=15)
        pdf.set_font("Arial", size=12)
        for index, row in df.iterrows():
            for col in df.columns:
                pdf.cell(40, 10, str(row[col]), 1)
            pdf.ln()
        pdf.output(output_path)

        return render_template('download_excel_pdf.html', filename=output_filename)
    else:
        return "Invalid file format. Please upload an Excel file."
    
    

@app.route('/uploaded', methods=['POST'])
def uploaded():
    # Check if file was submitted
    if 'file' not in request.files:
        return 'No file uploaded'

    file = request.files['file']

    # Check if the file is an image
    if file.filename == '':
        return 'No file selected'

    if file and allowed_file(file.filename):
        # Save the uploaded file to the upload folder
        file_path = os.path.join(app.config['UPLOAD_FOLDER'], file.filename)
        file.save(file_path)

        # Convert the image to Word document
        image = Image.open(file_path)
        document = Document()
        document.add_picture(file_path)
        word_file_path = os.path.join(app.config['UPLOAD_FOLDER'], f'{os.path.splitext(file.filename)[0]}.docx')
        document.save(word_file_path)

        return render_template('download_jpg_word.html', filename=os.path.basename(word_file_path))
    else:
        return 'Invalid file format'




@app.route('/download/<filename>', methods=['GET'])
def download(filename):
    return send_file(os.path.join(app.config['UPLOAD_FOLDER'], filename), as_attachment=True)


def allowed_file(filename):
    # Check if the file has an allowed extension
    allowed_extensions = {'jpg', 'jpeg', 'png','xlsx', 'xls','docx'}
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in allowed_extensions



if __name__ == '__main__':
    app.run(debug=True,host='0.0.0.0')
